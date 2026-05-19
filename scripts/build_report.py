"""
build_report.py — Step 5：将分析文字写入 Word 报告，保存到桌面

不依赖任何 AI API。分析内容由宿主 AI（Claude Code 对话）直接生成后，
以文本文件形式传入本脚本，脚本只负责排版和写 docx。

用法：
    # 模式 A：从分析文件 + 多个 note.json 生成报告
    python scripts/build_report.py --analysis analysis.txt note1.json note2.json ...

    # 模式 B：只用 note.json 生成"数据底稿"（无 AI 分析段落，只有互动数据表格）
    python scripts/build_report.py note1.json note2.json ...

分析文件格式（analysis.txt）：
    五个模块之间用 "==模块N==" 分隔：

    ==模块1==
    内容品类：干货卡片
    ...

    ==模块2==
    标题共性：...
    ...

    ==模块3==
    ...

    ==模块4==
    篇目 | 作者 | 点赞 | 收藏 | 评论 | 分享 | 藏赞比
    ...
    爆款原因分析：...

    ==模块5==
    已验证选题：...
    - 方向1
    ...

依赖：
    pip install python-docx

输出：
    ~/Desktop/小红书内容分析报告_YYYYMMDD_HHMM.docx
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

MODULE_SEP = re.compile(r"==模块(\d)==")

MODULE_TITLES = {
    "1": "模块 1：赛道定位卡",
    "2": "模块 2：结构共性提取",
    "3": "模块 3：观点共性提取",
    "4": "模块 4：互动数据对比面板",
    "5": "模块 5：选题矩阵",
}


# ──────────────────────────────────────────────
# 数据读取
# ──────────────────────────────────────────────

def load_notes(paths: list[str]) -> list[dict]:
    notes = []
    for p in paths:
        p = os.path.expanduser(p)
        if not os.path.isfile(p):
            raise FileNotFoundError(f"note.json 不存在：{p}")
        with open(p, "r", encoding="utf-8") as f:
            notes.append(json.load(f))
    return notes


def load_analysis(path: str) -> str:
    path = os.path.expanduser(path)
    if not os.path.isfile(path):
        raise FileNotFoundError(f"分析文件不存在：{path}")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# ──────────────────────────────────────────────
# 数据底稿：从 note.json 自动生成模块4表格
# ──────────────────────────────────────────────

def build_data_table(notes: list[dict]) -> str:
    """用 note.json 的互动数据自动生成模块4文本（无 AI 分析时的兜底）。"""
    lines = ["篇目 | 作者 | 点赞 | 收藏 | 评论 | 分享 | 藏赞比"]
    for i, note in enumerate(notes, 1):
        interact = note.get("interact", {})
        lines.append(
            f"{note.get('title', f'笔记{i}')} | "
            f"{note.get('author', '—')} | "
            f"{interact.get('liked', 0)} | "
            f"{interact.get('collected', 0)} | "
            f"{interact.get('commented', 0)} | "
            f"{interact.get('shared', 0)} | "
            f"{interact.get('collect_like_ratio', 0)}"
        )
    return "\n".join(lines)


# ──────────────────────────────────────────────
# Word 写入：模块4（表格行 + 爆款分析）
# ──────────────────────────────────────────────

def _write_module4_table(doc, body: str):
    """将模块4中的 | 分隔行写成真实 Word 表格，其余文字正常写段落。"""
    table_rows = []
    analysis_lines = []

    for line in body.splitlines():
        line = line.strip()
        if not line:
            continue
        if "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                table_rows.append(cells)
        else:
            analysis_lines.append(line)

    if table_rows:
        col_count = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=col_count)
        tbl.style = "Table Grid"
        for ri, row_data in enumerate(table_rows):
            for ci in range(col_count):
                cell_text = row_data[ci] if ci < len(row_data) else ""
                cell = tbl.cell(ri, ci)
                cell.text = cell_text
                for run in cell.paragraphs[0].runs:
                    if ri == 0:
                        run.bold = True
                    _set_font(run)
        doc.add_paragraph()

    for line in analysis_lines:
        if line.startswith("爆款原因分析"):
            _add_paragraph_with_font(doc, line, bold=True)
        else:
            _add_paragraph_with_font(doc, line)


# ──────────────────────────────────────────────
# Word 报告生成
# ──────────────────────────────────────────────

FONT_NAME = "宋体"


def _set_font(run):
    """将 run 的中西文字体统一设为宋体。"""
    from docx.oxml.ns import qn
    from lxml import etree
    run.font.name = FONT_NAME
    # 同时设置东亚文字（中文）字体
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)


def _add_paragraph_with_font(doc, text: str, style: str = None, bold: bool = False):
    """添加段落并对所有 run 应用宋体。"""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    _set_font(run)
    return p


def _apply_font_to_paragraph(paragraph):
    """对已有段落中的所有 run 应用宋体。"""
    for run in paragraph.runs:
        _set_font(run)


def build_docx(analysis_text: str, notes: list[dict], out_path: str):
    """
    将五模块分析文字写入 Word 文件。

    analysis_text: 含 ==模块N== 分隔符的完整分析文字
    notes:         note.json 列表（用于报告封面统计）
    out_path:      输出路径
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx 未安装。请运行：pip install python-docx")

    doc = Document()

    # 全局默认字体设为宋体
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)

    # 封面信息
    heading = doc.add_heading("小红书内容赛道分析报告", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font_to_paragraph(heading)

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    subtitle = _add_paragraph_with_font(doc, f"生成时间：{timestamp}　共 {len(notes)} 篇笔记")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 按分隔符切分五个模块
    # parts = ['前缀文字', '1', '模块1内容', '2', '模块2内容', ...]
    parts = MODULE_SEP.split(analysis_text)

    i = 1
    while i < len(parts) - 1:
        module_num = parts[i].strip()
        module_body = parts[i + 1].strip() if i + 1 < len(parts) else ""

        h = doc.add_heading(MODULE_TITLES.get(module_num, f"模块 {module_num}"), level=2)
        _apply_font_to_paragraph(h)

        if module_num == "4":
            _write_module4_table(doc, module_body)
        else:
            for line in module_body.splitlines():
                line = line.strip()
                if not line:
                    continue
                if line.startswith(("- ", "• ")):
                    _add_paragraph_with_font(doc, line[2:], style="List Bullet")
                elif line.endswith(("：", ":")):
                    _add_paragraph_with_font(doc, line, bold=True)
                else:
                    _add_paragraph_with_font(doc, line)

        i += 2

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    print(f"报告已保存：{out_path}")


# ──────────────────────────────────────────────
# 主函数
# ──────────────────────────────────────────────

def run(note_json_paths: list[str], analysis_file: str = "",
        account_name: str = "") -> str:
    """
    读取 note.json（及可选的分析文件）→ 写 Word 报告到桌面。

    analysis_file 为空时，仅写数据底稿（互动数据表格）。
    account_name  用于生成与仿写 skill 同名的报告文件，如 "粉领圆桌"
                  → 输出文件名：赛道分析_粉领圆桌_YYYYMMDD_HHMM.docx
                  省略时退回默认文件名：小红书内容分析报告_YYYYMMDD_HHMM.docx

    Returns:
        生成的 docx 路径
    """
    notes = load_notes(note_json_paths)
    print(f"读取了 {len(notes)} 篇笔记")

    if analysis_file:
        analysis_text = load_analysis(analysis_file)
        print(f"读取分析文件：{analysis_file}")
    else:
        # 无分析文件时，仅生成数据底稿（模块4表格）
        data_table = build_data_table(notes)
        analysis_text = f"==模块4==\n{data_table}"
        print("未提供分析文件，仅写入互动数据表格")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    if account_name:
        safe_name = re.sub(r"[^\w\u4e00-\u9fff\-]", "-", account_name.strip()).strip("-")
        filename = f"赛道分析_{safe_name}_{timestamp}.docx"
    else:
        filename = f"小红书内容分析报告_{timestamp}.docx"
    out_path = os.path.expanduser(f"~/Desktop/{filename}")

    build_docx(analysis_text, notes, out_path)
    return out_path


# ──────────────────────────────────────────────
# CLI 入口
# ──────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="将小红书笔记数据和 AI 分析文字写入 Word 报告（Step 5）"
    )
    parser.add_argument(
        "note_jsons", nargs="+",
        help="一个或多个 note.json 路径"
    )
    parser.add_argument(
        "--analysis", default="", dest="analysis_file",
        metavar="FILE",
        help="AI 分析文字文件（含 ==模块N== 分隔符）；省略则只写数据表格"
    )
    parser.add_argument(
        "--account-name", default="", dest="account_name",
        metavar="NAME",
        help="账号名称，用于生成与仿写 skill 同名的报告文件（如 '粉领圆桌'）"
    )
    args = parser.parse_args()

    try:
        run(args.note_jsons, analysis_file=args.analysis_file,
            account_name=args.account_name)
    except Exception as e:
        print(f"错误：{e}", file=sys.stderr)
        sys.exit(1)
